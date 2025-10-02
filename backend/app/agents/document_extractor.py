import io
from typing import Optional
def extract_text_from_file(file_content: bytes, filename: str) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        return _extract_from_pdf(file_content)
    elif filename_lower.endswith(('.doc', '.docx')):
        return _extract_from_word(file_content)
    elif filename_lower.endswith('.txt'):
        return file_content.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: {filename}. Supported types: PDF, Word (.doc/.docx), TXT")
def _extract_from_pdf(file_content: bytes) -> str:
    try:
        import PyPDF2
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                print(f"Warning: Failed to extract text from page {page_num + 1}: {e}")
                continue
        if not text_parts:
            raise ValueError("No text could be extracted from PDF")
        return "\n\n".join(text_parts)
    except ImportError:
        raise ValueError("PyPDF2 library is required for PDF extraction")
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")
def _extract_from_word(file_content: bytes) -> str:
    try:
        from docx import Document
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        if not text_parts:
            raise ValueError("No text could be extracted from Word document")
        return "\n\n".join(text_parts)
    except ImportError:
        raise ValueError("python-docx library is required for Word document extraction")
    except Exception as e:
        raise ValueError(f"Failed to extract text from Word document: {str(e)}")
def validate_file_type(filename: str) -> bool:
    filename_lower = filename.lower()
    return filename_lower.endswith(('.pdf', '.doc', '.docx', '.txt'))
def get_supported_extensions() -> list:
    return ['.pdf', '.doc', '.docx', '.txt']
